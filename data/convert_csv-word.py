import csv
from docx import Document
from pathlib import Path
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

def keep_with_next(paragraph):
    # 段落を次の段落と同じページにする設定
    pPr = paragraph._element.get_or_add_pPr()
    keep_next = OxmlElement('w:keepNext')
    pPr.append(keep_next)

def csv_to_docx(csv_file, docx_file, title):
    doc = Document()

     # タイトル（中央揃え・黒文字）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)   # 黒
    run.font.name = "Meiryo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")

    doc.add_paragraph("") # タイトル後の空行

    label_map = {
                "name": "名称",
                "timing": "タイミング",
                "target": "対象",
                "condition": "条件",
                "effect": "効果",
                "description": "解説",
            }

    with open(csv_file, encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            # マギの名称
            p = doc.add_paragraph("")
            run = p.add_run(f"《{row['name']}》")
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Meiryo"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")
            keep_with_next(p) #ページ跨ぎ防止
            
            # タイミング～解説
            for key in ["timing","target","condition","effect","description"]:
                value = row.get(key.lower(), "")
                p = doc.add_paragraph(f"{label_map[key]}：{value}")
                p.style.font.name = "Meiryo"
                p.style._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")
                p.style.font.size = Pt(9) #フォントサイズ
                p.paragraph_format.line_spacing = 0.9     # 行間（段落内）
                p.paragraph_format.space_before = Pt(7)   # 段落前の余白
                p.paragraph_format.space_after = Pt(0.5)  # 段落後の余白（半ポイント）
                
                if key != "description":
                    keep_with_next(p) #ページ跨ぎ防止（最後以外）

            doc.add_paragraph("") # マギごとの間隔

    doc.save(docx_file)

def convert_and_log(csv_file, docx_file, title):
    csv_to_docx(csv_file, docx_file, title)
    print(f"[OK] {Path(csv_file).name:<15} → {Path(docx_file).name}")

# 変換したいファイルをここに追加する
files_to_convert = [
    # ("CSV名","word名","word内のタイトル名")
    ("magi_common.csv", "共通マギ_一覧.docx", "共通マギ一覧"),
    ("magi_pc.csv",     "PC用マギ_一覧.docx", "PC用マギ一覧"),
    ("magi_clan.csv",   "クラン用マギ_一覧.docx", "クラン用マギ一覧"),
]

def main():
    for src, dst, title in files_to_convert:
        convert_and_log(src, dst, title)
    print(f"=== {len(files_to_convert)} files converted successfully ===")

if __name__ == "__main__":
    main()