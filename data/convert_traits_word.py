import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ------------- Word 小技（XML直叩き） ----------------
def keep_with_next(paragraph):
    """次の段落と同じページに保つ（Wordの[次の段落と一緒にする]）"""
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))

def keep_lines_together(paragraph):
    """段落内で行を分割しない（Wordの[段落内で改ページしない]）"""
    pPr = paragraph._element.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepLines'))

def apply_body_style(p):
    """ご指定の段落スタイルを適用"""
    # 段落の既定スタイルに直接フォント指定を入れるのは難があるので run で設定
    # ただし python-docx では paragraph.style.font は効きづらいので、ここでは段落の行間・余白だけ確実に適用
    p.paragraph_format.line_spacing = 0.9        # 行間（段落内）
    p.paragraph_format.space_before = Pt(7)      # 段落前
    p.paragraph_format.space_after = Pt(0.5)     # 段落後

def add_run_meiryo(p, text, size_pt=9, bold=False):
    """Meiryo 指定の run を追加"""
    r = p.add_run(text)
    r.font.name = "Meiryo"                       # 西欧
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Meiryo")  # 東アジア
    r.font.size = Pt(size_pt)
    r.bold = bold
    return r


# ------------- メイン処理 ----------------
def csv_to_traits_docx(csv_file: str, docx_file: str):
    doc = Document()

    # タイトル
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_meiryo(p_title, "特性の例一覧", size_pt=14, bold=True)
    # タイトル後の少し余白
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

    # CSV 読み込み
    with open(csv_file, encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        current_cat = None
        first_block = True

        for row in reader:
            cat = (row.get("category") or "").strip()
            name = (row.get("name") or "").strip()
            desc = (row.get("description") or "").strip()

            # カテゴリが変わったら改ページ
            if current_cat is None:
                current_cat = cat
            elif cat != current_cat:
                doc.add_page_break()
                current_cat = cat
                first_block = True  # ページ先頭フラグ更新

            # ページ先頭にカテゴリ見出しを入れたい場合（必要なければこの2行を削除）
            if first_block:
                p_cat = doc.add_paragraph()
                apply_body_style(p_cat)
                add_run_meiryo(p_cat, f"【{cat}】", size_pt=12, bold=True)
                keep_lines_together(p_cat)
                first_block = False

            # --- 特性名 ---
            p_name = doc.add_paragraph()
            apply_body_style(p_name)  # 行間・余白はベース設定
            add_run_meiryo(p_name, f"《{name}》", size_pt=10, bold=True)
            p_name.paragraph_format.space_after = Pt(0)   # 説明との間を小さく
            keep_with_next(p_name)
            keep_lines_together(p_name)

            # --- 説明文 ---
            p_desc = doc.add_paragraph()
            apply_body_style(p_desc)
            add_run_meiryo(p_desc, desc, size_pt=9, bold=False)
            p_desc.paragraph_format.space_before = Pt(0)  #段落間の余白制御
            p_desc.paragraph_format.space_after = Pt(11)  #特性同士の間隔を広めに
            keep_lines_together(p_desc)


    doc.save(docx_file)


# ------------- 使い方例 ----------------
if __name__ == "__main__":
    csv_to_traits_docx("traits.csv", "特性の例一覧.docx")
    print("[OK] traits.csv           → 特性の例一覧.docx")
